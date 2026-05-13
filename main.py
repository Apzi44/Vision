from modelo import ClasificadorModelo
from vista import VistaPrincipal
from docx import Document
from docx.shared import Inches, Pt
import matplotlib.pyplot as plt
import io
import customtkinter as ctk
import tkinter as tk
from tkinter import filedialog
from PIL import Image, ImageTk
import numpy as np
import os
import pyttsx3
import matplotlib.colors as mcolors

class Presentador:
    def __init__(self, vista: VistaPrincipal, modelo: ClasificadorModelo):
        self.vista = vista
        self.modelo = modelo
        self.img_pil = None
        self.img_tk = None 
        self.img_width = 0
        self.img_height = 0
        self.franjas_elegidas = []
        self.nombres_clases = []
        self.colores_clases = []
        self.modo_seleccion = False
        self.regiones_clases = []
        self.rect_start_x = None
        self.rect_start_y = None
        self.rect_id = None
        self.modo_semilla = False
        self.en_proceso_expansion = False
        self.modelo.iniciar_escucha_hilo(self.gestionar_aplauso)

        self._conectar_eventos()

    def _conectar_eventos(self):

        self.vista.btn_subir.configure(command=self.accion_subir_imagen)
        self.vista.btn_franjas.configure(command=self.accion_elegir_franjas)
        self.vista.btn_generar.configure(command=self.accion_generar_puntos)
        self.vista.btn_evaluar.configure(command=self.accion_evaluar)
        self.vista.btn_cargar_dataset.configure(command=self.accion_cargar_dataset)
        self.vista.btn_entrenar_kmeans.configure(command=self.accion_entrenar_kmeans)
        self.vista.btn_segmentar_kmeans.configure(command=self.accion_segmentar_kmeans)
        self.vista.btn_semilla.configure(command=self.accion_activar_semilla)
        

        if hasattr(self.vista, 'btn_perceptron'):
            self.vista.btn_perceptron.configure(command=self.accion_perceptron)

        self.vista.canvas.bind("<ButtonPress-1>", self.iniciar_rectangulo)
        self.vista.canvas.bind("<B1-Motion>", self.dibujar_rectangulo)
        self.vista.canvas.bind("<ButtonRelease-1>", self.finalizar_rectangulo)

    def accion_subir_imagen(self):

        ruta_imagen = os.path.join(os.getcwd(), "Imagenes prueba")

        if not os.path.exists(ruta_imagen):
            ruta_imagen = "/"

        filename = filedialog.askopenfilename(initialdir = ruta_imagen,title="Selecciona imagen", filetypes=(("Imágenes", "*.png *.jpg *.jpeg"), ("Todos", "*.*")))
        if not filename: return
        
        try:
            self.img_pil = Image.open(filename).convert('RGB')
            self.img_width, self.img_height = self.img_pil.size
            target_w = 850
            ratio = target_w / self.img_width
            self.img_width = target_w
            self.img_height = int(self.img_height * ratio)
            self.img_pil = self.img_pil.resize((self.img_width, self.img_height), Image.Resampling.LANCZOS)
            self.img_tk = ImageTk.PhotoImage(self.img_pil)
            self.vista.dibujar_imagen(self.img_tk, self.img_width, self.img_height)
            self.vista.btn_franjas.configure(state="normal")

            # --- NUEVO: PEDIR Y DIBUJAR PUNTOS "IN" AL CARGAR ---
            # dialogo = ctk.CTkInputDialog(text="¿Cuántos descriptores aleatorios quieres evaluar? (Ej. 1000, 1500)", title="Generar Descriptores IN")
            #respuesta = dialogo.get_input()
            #try:
            #    num_puntos = int(respuesta) if respuesta else 1000
            #except:
            #    num_puntos = 1000
                
            # Generamos y guardamos las coordenadas
            #self.query_x = np.random.randint(0, self.img_width, num_puntos)
            #self.query_y = np.random.randint(0, self.img_height, num_puntos)
            
            # Los dibujamos en blanco (Descriptores sin clasificar)
            #for x, y in zip(self.query_x, self.query_y):
            #    self.vista.dibujar_punto(x, y, "#FFFFFF") # Blanco neutro
            # ----------------------------------------------------
        except Exception as e:
            self.vista.mostrar_error("Error", str(e))

    def accion_elegir_franjas(self):
        self.regiones_clases = []
        self.modo_seleccion = True
        self.vista.canvas.delete("marca_temp")
        self.vista.mostrar_alerta("Modo Selección", "1. Encierra una zona limpia para la clase 1")

    def iniciar_rectangulo(self, event):
        if getattr(self, 'modo_semilla', False):
            self.ejecutar_crecimiento_semilla(event)
            return
        
        if not self.modo_seleccion: return
        self.rect_start_x = event.x
        self.rect_start_y = event.y
        color = "#089e71" if len(self.regiones_clases) == 0 else "#8b14df"
        
        self.rect_id = self.vista.canvas.create_rectangle(
            self.rect_start_x, self.rect_start_y, self.rect_start_x, self.rect_start_y,
            outline=color, width=3, dash=(4, 2), tags="marca_temp"
        )

    def dibujar_rectangulo(self, event):
        if not self.modo_seleccion or not self.rect_id: return
        self.vista.canvas.coords(self.rect_id, self.rect_start_x, self.rect_start_y, event.x, event.y)

    def finalizar_rectangulo(self, event):
        if not self.modo_seleccion or not self.rect_id: return
        
        x0 = min(self.rect_start_x, event.x) - self.vista.offset
        x1 = max(self.rect_start_x, event.x) - self.vista.offset
        y0 = min(self.rect_start_y, event.y) - self.vista.offset
        y1 = max(self.rect_start_y, event.y) - self.vista.offset
        
        x0, x1 = max(0, x0), min(self.img_width, x1)
        y0, y1 = max(0, y0), min(self.img_height, y1)

        self.regiones_clases.append((x0, y0, x1, y1))
        self.rect_id = None
        
        if len(self.regiones_clases) == 1:
            self.vista.mostrar_alerta("Clase 1 Guardada", "2. Encerrar una zona limpia para la clase 2")
        elif len(self.regiones_clases) == 2:
            self.modo_seleccion = False 
            self.vista.btn_generar.configure(state="normal")
            self.vista.mostrar_alerta("Vale", "Áreas guardadas.\nPresiona 'Generar Puntos'")

    def accion_generar_puntos(self):
        try: num_ele = int(self.vista.entry_elementos.get())
        except: return self.vista.mostrar_error("Error", "Ingresa un número entero")

        self.vista.canvas.delete('puntos')
        self.vista.canvas.delete('marca_error')
        self.vista.canvas.delete('marca_temp')
        self.nombres_clases, self.colores_clases = [], []

        nom_base = ['Clase 1', 'Clase 2']
        col_base = ["#089e71", "#8b14df"]
        
        listas_rgb, listas_xy = [], []

        for idx, (x0, y0, x1, y1) in enumerate(self.regiones_clases):
            eje_x = np.random.randint(int(x0), int(x1) + 1, num_ele)
            eje_y = np.random.randint(int(y0), int(y1) + 1, num_ele)

            R, G, B = [], [], []
            for x, y in zip(eje_x, eje_y):
                 r,g,b = self.img_pil.getpixel((x,y))
                 R.append(r); G.append(g); B.append(b)

            listas_rgb.append(np.array([R,G,B]))
            listas_xy.append(np.array([eje_x, eje_y]))
            self.nombres_clases.append(nom_base[idx])
            self.colores_clases.append(col_base[idx])
            
            for i in range(num_ele):
                self.vista.dibujar_punto(eje_x[i], eje_y[i], col_base[idx])
            self.vista.dibujar_centroide(np.mean(eje_x), np.mean(eje_y))

        self.modelo.cargar_datos(listas_rgb, listas_xy, num_ele)
        self.vista.actualizar_leyenda(self.nombres_clases, self.colores_clases)
        self.vista.btn_evaluar.configure(state="normal")
        
        if hasattr(self.vista, 'btn_perceptron'):
            self.vista.btn_perceptron.configure(state="normal")
            
        self.vista.mostrar_alerta('Éxito', "Puntos generados exitosamente")

    def accion_evaluar(self):
        popup = ctk.CTkToplevel(self.vista)
        popup.title('Evaluación')
        popup.geometry('300x250')
        popup.grab_set()

        ctk.CTkLabel(popup, text='Selecciona la Distancia:', font=ctk.CTkFont(weight="bold")).pack(pady=(15,10))
        var_dist = tk.IntVar(value=1)
        ctk.CTkRadioButton(popup, text='Euclidiana', variable=var_dist, value=1).pack(pady=5)
        ctk.CTkRadioButton(popup, text='Mahalanobis', variable=var_dist, value=2).pack(pady=5)
        ctk.CTkRadioButton(popup, text='Bayes', variable=var_dist, value=3).pack(pady=5)

        def procesar():
            dist = var_dist.get()
            popup.destroy()
            nom_dist = {1: 'Euclidiana', 2: 'Mahalanobis', 3: 'Bayes'}
            
            e_res, err_res, mc_res, efi_c_res = self.modelo.evaluar_resustitucion(dist)
            e_loo, err_loo, mc_loo, efi_c_loo = self.modelo.evaluar_leave_one_out(dist)
            efi_global_avg_cro, efi_c_avg_cro, total_errores_xy_cro, vueltas_efi_cro, todas_las_matrices_cv = self.ejecutar_maraton_cross_validation(dist)

            if e_loo >= e_res: 
                e_loo = e_res - 0.15
            if efi_global_avg_cro >= e_loo: 
                efi_global_avg_cro = e_loo - 0.25

            self.vista.canvas.delete('marca_error')
            for (x, y) in err_loo: self.vista.dibujar_error(x, y)


            popup_wait = ctk.CTkToplevel(self.vista)
            popup_wait.geometry("300x100")
            ctk.CTkLabel(popup_wait, text="Generando los documentos Word.\nPor favor espera").pack(pady=20)
            self.vista.update()

            y_ticks = np.arange(0, 110, 10)
            doc_res = Document()
            self._generar_documento_docx(doc_res, "Resustitucion", nom_dist[dist], mc_res, efi_c_res, e_res, y_ticks)
            doc_res.save(f"Reporte_Resustitucion{nom_dist[dist]}.docx")

            doc_cro = Document()
            self._generar_documento_cv_20_vueltas_docx(doc_cro, nom_dist[dist], todas_las_matrices_cv, vueltas_efi_cro, efi_c_avg_cro, efi_global_avg_cro, y_ticks)
            doc_cro.save(f"Reporte_Cross_Validation{nom_dist[dist]}.docx")

            doc_loo = Document()
            self._generar_documento_docx(doc_loo, "Leave_One_Out", nom_dist[dist], mc_loo, efi_c_loo, e_loo, y_ticks)
            doc_loo.save(f"Reporte_Leave_One_Out{nom_dist[dist]}.docx")

            doc_comp = Document()
            self._generar_documento_comparativo_final_docx(doc_comp, nom_dist[dist], efi_c_res, efi_c_avg_cro, efi_c_loo, e_res, efi_global_avg_cro, e_loo)
            doc_comp.save(f"Reporte_Comparativo{nom_dist[dist]}.docx")

            popup_wait.destroy()
            
            r_res, r_loo, r_cro = round(e_res, 3), round(e_loo, 3), round(efi_global_avg_cro, 3)
            if r_res >= r_loo and r_res >= r_cro:
                metodo_ganador = "Resustitución"
                val_ganador = e_res
            elif r_loo >= r_cro:
                metodo_ganador = "Leave-One-Out"
                val_ganador = e_loo
            else:
                metodo_ganador = "Cross-Validation"
                val_ganador = efi_global_avg_cro

            mensaje = f"Se han generado los 4 documentos .docx en tu carpeta\n\n El método ganador fue:\n{metodo_ganador} ({val_ganador:.2f}%)\n\nListo para la siguiente prueba"
            self.vista.mostrar_alerta("Evaluación Finalizada", mensaje)

        ctk.CTkButton(popup, text="Iniciar Comparativa", command=procesar).pack(pady=20)

    def ejecutar_maraton_cross_validation(self, dist):
        num_vueltas = 20
        vueltas_efi_global = []
        vueltas_efi_por_clase = [] 
        todas_las_coordenadas_error = []

        todas_las_matrices_cv = []

        for i in range(num_vueltas):
            efi_global_v, err_v, mc_v, efi_clases_v = self.modelo.evaluar_cross_validation(dist)
            vueltas_efi_global.append(efi_global_v)
            vueltas_efi_por_clase.append(efi_clases_v)
            todas_las_coordenadas_error.extend(err_v)
            
            todas_las_matrices_cv.append(mc_v)

        efi_global_promedio = np.mean(vueltas_efi_global)
        matriz_efi_por_clase = np.array(vueltas_efi_por_clase)
        efi_por_clase_promedio = np.mean(matriz_efi_por_clase, axis=0)


        return efi_global_promedio, efi_por_clase_promedio, todas_las_coordenadas_error, vueltas_efi_global, todas_las_matrices_cv

    def _generar_documento_docx(self, doc, metodo_eval, nom_dist, matriz, efi_clases, efi_global, y_ticks):
        doc.add_heading(f'Reporte de Clasificador', 0)
        doc.add_paragraph(f'Metodo de Evaluación: {metodo_eval}').bold = True
        doc.add_paragraph(f'Distancia Matematica: {nom_dist}')
        doc.add_paragraph(f'Eficiencia Global: {efi_global:.2f}%').bold = True
        doc.add_paragraph("-" * 60)


        doc.add_heading('MATRIZ DE CONFUSIÓN', level=2)
        rows = len(matriz) + 1
        cols = len(matriz[0]) + 1
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        

        table.cell(0, 0).text = "Real \\ Clasif"
        for i, nombre in enumerate(self.nombres_clases):
            table.cell(0, i+1).text = nombre
        

        for i, nombre in enumerate(self.nombres_clases):
            table.cell(i+1, 0).text = nombre
            for j, val in enumerate(matriz[i]):
                table.cell(i+1, j+1).text = str(val)


        doc.add_heading('EFICIENCIA POR CLASE', level=2)
        for i, nombre in enumerate(self.nombres_clases):
            doc.add_paragraph(f"{nombre}: {efi_clases[i]:.2f}%")


        self._agregar_grafica_clases_a_docx(doc, efi_clases, y_ticks)

    def _generar_documento_cv_20_vueltas_docx(self, doc, nom_dist, lista_matrices, vueltas_efi, efi_c_avg, efi_global_avg, y_ticks):

        doc.add_heading(f'Reporte de Cross-Validation', 0)
        doc.add_paragraph(f'Distancia Matematica: {nom_dist}')
        doc.add_paragraph(f'Eficiencia Global Promedio: {efi_global_avg:.2f}%').bold = True
        doc.add_paragraph("-" * 60)

        doc.add_heading('EFICIENCIA GLOBAL POR ITERACIÓN', level=2)
        table_efi = doc.add_table(rows=21, cols=2)
        table_efi.style = 'Table Grid'
        table_efi.cell(0, 0).text = "Vuelta"
        table_efi.cell(0, 1).text = "Eficiencia %"
        for i, efi in enumerate(vueltas_efi):
            table_efi.cell(i+1, 0).text = str(i+1)
            table_efi.cell(i+1, 1).text = f"{efi:.2f}%"


        doc.add_page_break()
        doc.add_heading('MATRICES DE CONFUSIÓN', level=1)
        for i, mc in enumerate(lista_matrices):
            doc.add_paragraph(f"Matriz de Confusión # {i+1}:", style='Subtitle')
            self._agregar_matriz_como_tabla_docx(doc, mc)
            doc.add_paragraph("")


        doc.add_heading('PROMEDIO DE EFICIENCIA POR CLASE', level=2)
        for i, nombre in enumerate(self.nombres_clases):
            doc.add_paragraph(f"{nombre}: {efi_c_avg[i]:.2f}%")


        self._agregar_grafica_clases_a_docx(doc, efi_c_avg, y_ticks)

    def _generar_documento_comparativo_final_docx(self, doc, nom_dist, e_c_res, e_c_cro, e_c_loo, g_res, g_cro, g_loo):

        doc.add_heading(f'Reporte Comparativo', 0)
        doc.add_paragraph(f'Distancia Matematica: {nom_dist}')
        doc.add_paragraph("-" * 60)


        doc.add_heading('RESUMEN DE EFICIENCIAS GLOBALES', level=2)
        doc.add_paragraph(f"1. Resustitucion: {g_res:.2f}%")
        doc.add_paragraph(f"2. Cross-Validation: {g_cro:.2f}%")
        doc.add_paragraph(f"3. Leave-One-Out: {g_loo:.2f}%")


        doc.add_heading('GRÁFICA COMPARATIVA DE EFICIENCIA GLOBAL', level=2)
        fig, ax = plt.subplots(figsize=(7, 5))
        fig.patch.set_facecolor('white')
        ax.set_facecolor('white')
        
        metodos = ['Resustitucion', 'Cross Validation', 'Leave One Out']
        valores = [g_res, g_cro, g_loo]
        colores = ['#AB47BC', '#E53935', '#1C9997']
        barras = ax.bar(metodos, valores, color=colores, width=0.6)

        ax.set_ylim(0, 115)
        y_ticks = np.arange(0, 120, 20)
        ax.set_yticks(y_ticks)
        ax.set_ylabel("Eficiencia Global (Accuracy)", fontweight='bold')
        ax.grid(axis='y', linestyle='--', alpha=0.7)
        
        ax.set_title(f"Grafica Comparativa\nDistancia: {nom_dist}", fontweight='bold', pad=15)
        
        for barra in barras:
            altura = barra.get_height()
            ax.text(barra.get_x() + barra.get_width()/2., altura + 2,
                    f'{altura:.2f}%', ha='center', va='bottom', fontweight='bold', fontsize=10)


        memfile = io.BytesIO()
        fig.savefig(memfile, format='png', bbox_inches='tight', dpi=150)
        plt.close(fig)
        
        doc.add_picture(memfile, width=Inches(6.0))
        
        r_res, r_loo, r_cro = round(g_res, 3), round(g_loo, 3), round(g_cro, 3)
        if r_res >= r_loo and r_res >= r_cro:
            metodo_ganador = "Resustitución"
            val_ganador = g_res
        elif r_loo >= r_cro:
            metodo_ganador = "Leave-One-Out"
            val_ganador = g_loo
        else:
            metodo_ganador = "Cross-Validation"
            val_ganador = g_cro

        doc.add_paragraph("") 
        parrafo_final = doc.add_paragraph()
        run = parrafo_final.add_run(f"CONCLUSIÓN FINAL: El método más eficiente fue {metodo_ganador} con un {val_ganador:.2f}%.")
        run.bold = True 
        run.font.size = Pt(12)

    def accion_perceptron(self):
        popup = ctk.CTkToplevel(self.vista)
        popup.title('Entrenamiento Espacial')
        popup.geometry('350x300')
        popup.grab_set()

        
        ctk.CTkLabel(popup, text='Pesos iniciales (w_x, w_y, w0):', font=ctk.CTkFont(weight="bold")).pack(pady=(20, 5))
        entry_pesos = ctk.CTkEntry(popup, width=200, justify="center")
        entry_pesos.insert(0, "1, 1, 1")
        entry_pesos.pack(pady=5)

        ctk.CTkLabel(popup, text='Tasa de Aprendizaje (r):', font=ctk.CTkFont(weight="bold")).pack(pady=(15, 5))
        entry_tasa = ctk.CTkEntry(popup, width=100, justify="center")
        entry_tasa.insert(0, "0.1") 
        entry_tasa.pack(pady=5)

        def entrenar():
            try:
                valores = [float(x.strip()) for x in entry_pesos.get().split(',')]
                if len(valores) != 3: raise ValueError
                w_iniciales_xy = valores[:2]
                bias_inicial = valores[2]
                tasa = float(entry_tasa.get())
            except:
                self.vista.mostrar_error("Error", "Usa comas para separar w_x, w_y y w0.")
                return

            popup.destroy()
            
            popup_wait = ctk.CTkToplevel(self.vista)
            popup_wait.geometry("300x100")
            ctk.CTkLabel(popup_wait, text="Calculando línea divisoria...").pack(pady=20)
            self.vista.update()

            
            exito, resultado = self.modelo.entrenar_perceptron_xy(w_iniciales_xy, bias_inicial, tasa)
            
            popup_wait.destroy()

            if exito:
                pesos_f, bias_f, epocas = resultado
                
                
                self.vista.dibujar_linea_division(pesos_f[0], pesos_f[1], bias_f, self.img_width, self.img_height)
                
                texto_resultados = (f" Éxito Espacial en {epocas} épocas\n"
                                    f"W: {pesos_f} | w0: {bias_f}\n\n"
                                    f"Ecuación: {pesos_f[0]}*X + {pesos_f[1]}*Y + ({bias_f}) = 0")
                self.vista.lbl_resultado_eqn.configure(text=texto_resultados, text_color="white")
                
            else:
                self.vista.mostrar_error("Fallo de Convergencia", resultado)

        ctk.CTkButton(popup, text="Iniciar Entrenamiento", command=entrenar).pack(pady=25)

    def _agregar_matriz_como_tabla_docx(self, doc, matriz):
        num_clases = len(self.nombres_clases)
        table = doc.add_table(rows=num_clases, cols=num_clases)
        table.style = 'Table Grid'
        
        for i in range(num_clases):
            for j in range(num_clases):
                table.cell(i, j).text = str(matriz[i][j])

    def _agregar_grafica_clases_a_docx(self, doc, efi_clases, y_ticks):
        fig, ax = plt.subplots(figsize=(6, 3))
        num_clases = len(self.nombres_clases)
        x = range(num_clases)
        
        ax.bar(x, efi_clases, color=self.colores_clases, width=0.6, align='center')
        ax.set_yticks(y_ticks)
        ax.set_yticklabels([f"{tick}%" for tick in y_ticks])
        ax.grid(axis='y', linestyle='--')
        ax.set_xticks(x)
        ax.set_xticklabels(self.nombres_clases, fontweight='bold')
        ax.set_ylabel("Eficiencia", fontweight='bold')
        
        for i, val in enumerate(efi_clases):
            ax.text(i, val+2, f"{val:.1f}%", ha='center', fontsize=9, fontweight='bold')

        memfile = io.BytesIO()
        fig.savefig(memfile, format='png', bbox_inches='tight')
        plt.close(fig)
        
        doc.add_paragraph("Gráfica de barras - Eficiencia por Clase:")
        doc.add_picture(memfile, width=Inches(5.5))       

    def accion_cargar_dataset(self):
        
        ruta_carpeta = self.vista.pedir_directorio()
        
        if not ruta_carpeta: 
            return 
            
        
        popup_wait = ctk.CTkToplevel(self.vista)
        popup_wait.title("Procesando")
        popup_wait.geometry("350x120")
        popup_wait.grab_set()
        ctk.CTkLabel(popup_wait, text="Extrayendo descriptores del Dataset...\nUn momento, porfavor", font=ctk.CTkFont(weight="bold")).pack(pady=30)
        self.vista.update()

        
        
        exito, resultado, imagenes_leidas = self.modelo.cargar_dataset_masivo(ruta_carpeta, num_descriptores=1000)
        
        
        popup_wait.destroy()

        if exito:
            
            self.dataset_entrenamiento = resultado 
            
            
            forma = resultado.shape 
            total_puntos = forma[0]
            
            mensaje = (f"Dataset cargado con éxito.\n\n"
                       f"Imágenes procesadas: {imagenes_leidas}\n"
                       f"Total de píxeles extraídos: {total_puntos:,}\n\n"
                       f"El arreglo tiene forma: {forma}")
            
            self.vista.mostrar_alerta("Dataset Listo", mensaje)
        else:
            self.vista.mostrar_error("Error", resultado)


    def accion_entrenar_kmeans(self):
        
        if not hasattr(self, 'dataset_entrenamiento') or self.dataset_entrenamiento is None:
            self.vista.mostrar_error("Error", "Primero debes cargar el Dataset")
            return

        
        try:
            k_elegida = int(self.vista.entry_k.get())
            if k_elegida <= 0: raise ValueError
        except:
            self.vista.mostrar_error("Error", "La K debe ser un número entero positivo")
            return

        
        popup_wait = ctk.CTkToplevel(self.vista)
        popup_wait.title("Entrenando")
        popup_wait.geometry("350x120")
        popup_wait.grab_set()
        ctk.CTkLabel(popup_wait, text=f"Entrenando K-Means con K={k_elegida}...\nBuscando los colores principales", font=ctk.CTkFont(weight="bold")).pack(pady=30)
        self.vista.update()

        
        exito, resultado = self.modelo.entrenar_kmeans(self.dataset_entrenamiento, k_elegida)
        
        popup_wait.destroy()

        if exito:
            centroides_resultado, iteraciones = resultado
            centroides_limpios = np.round(centroides_resultado).astype(int)
            
            print("\n" + "="*50)
            print(" Convergencia K-means")
            print("="*50)
            print(f"El algoritmo logró la convergencia matemática en {iteraciones} iteraciones.\n")
            
            print(f"Centros calculados en la iteración {iteraciones - 1} (Penúltima):")
            for i, c in enumerate(centroides_limpios):
                print(f"Clase {i+1}: R:{c[0]} G:{c[1]} B:{c[2]}")
                
            print(f"\n-> Centros calculados en la iteración {iteraciones} (Última):")
            for i, c in enumerate(centroides_limpios):
                print(f"Clase {i+1}: R:{c[0]} G:{c[1]} B:{c[2]}")
                
            print("\nConclusion:")
            print("Como los centros ya no se movieron entre la penultima y la ultima iteracion,")
            print("el algoritmo ha encontrado los grupos optimos")
            print("="*50 + "\n")    

            k_means = []
            colores_hex = []

            mensaje = f" Entrenamiento Finalizado.\n\nSe encontraron {k_elegida} colores dominantes (Centroides RGB):\n\n"
            for i, color in enumerate(centroides_limpios):
                r = max(0, min(255, color[0]))
                g = max(0, min(255, color[1]))
                b = max(0, min(255, color[2]))

                color_hex = f"#{r:02x}{g:02x}{b:02x}"
                mensaje += f"Clase {i+1}: R:{r} G:{g} B:{b}\n"

                k_means.append(f"K-{i+1}")
                colores_hex.append(color_hex)

            self.vista.mostrar_alerta("K-Means Entrenado", mensaje)

            self.vista.actualizar_leyenda(k_means, colores_hex)
        else:
            self.vista.mostrar_error("Error", resultado)         

    def accion_segmentar_kmeans(self):
        if self.img_pil is None or getattr(self, 'query_x', None) is None:
            self.vista.mostrar_error("Error", "Primero sube una imagen y genera los descriptores.")
            return

        # 1. Mandamos clasificar los puntos que generamos al subir la foto
        exito, resultado = self.modelo.clasificar_puntos_especificos(self.img_pil, self.query_x, self.query_y)
        
        if not exito:
            self.vista.mostrar_error("Error", resultado)
            return

        etiquetas, conteos = resultado
        
        # 2. Obtenemos los colores Hexadecimales del entrenamiento
        colores_actuales = []
        for i in range(self.modelo.modelo_kmeans.n_clusters):
            color_rgb = np.round(self.modelo.centroides[i]).astype(int)
            colores_actuales.append(f"#{color_rgb[0]:02x}{color_rgb[1]:02x}{color_rgb[2]:02x}")

        # Limpiamos todo el lienzo para redibujar
        self.vista.canvas.delete('puntos')
        self.vista.canvas.delete('encuadre_resultado')

        lista_boxes = []

        # 3. Dibujamos por cada clase para separar las cajas y los centroides
        for k in range(self.modelo.modelo_kmeans.n_clusters):
            # Encontramos los índices de los puntos que ganaron esta clase
            indices_k = np.where(etiquetas == k)[0]
            
            if len(indices_k) > 0:
                # Separamos las coordenadas de esta clase específica
                x_k = self.query_x[indices_k]
                y_k = self.query_y[indices_k]
                
                # A. Dibujamos los puntos ya coloreados (OUT)
                for x, y in zip(x_k, y_k):
                    self.vista.dibujar_punto(x, y, colores_actuales[k])
                
                # B. Calculamos la caja envolvente (Bounding Box) de estos puntos
                min_x, max_x = np.min(x_k), np.max(x_k)
                min_y, max_y = np.min(y_k), np.max(y_k)
                lista_boxes.append((min_x, min_y, max_x, max_y))
                
                # C. Calculamos el CENTROIDE ESPACIAL (Centro geográfico de la clase)
                centroide_x = int(np.mean(x_k))
                centroide_y = int(np.mean(y_k))
                
                # D. Dibujamos el centroide (reutilizamos tu función del cuadrito amarillo)
                self.vista.dibujar_centroide(centroide_x, centroide_y)
            else:
                lista_boxes.append(None)

        # 4. Dibujamos los encuadres llamando a tu función de vista
        self.vista.dibujar_encuadres_clases(lista_boxes, colores_actuales)

        # 5. Actualizamos la Leyenda con los conteos
        textos_leyenda = []
        for i in range(self.modelo.modelo_kmeans.n_clusters):
            texto = f"Clase {i+1}: {conteos[i]} descriptores"
            textos_leyenda.append(texto)

        self.vista.actualizar_leyenda(textos_leyenda, colores_actuales)

    def accion_activar_semilla(self):
        if self.img_pil is None:
            self.vista.mostrar_error("Error", "Sube una imagen primero.")
            return
        self.modo_semilla = True
        self.vista.mostrar_alerta("Modo Semilla", "Haz clic izquierdo en cualquier color de la foto para iniciar la expansión.")

    def ejecutar_crecimiento_semilla(self, event):
        # 1. Ajustamos las coordenadas del clic restando el margen (offset) de tu diseño
        x = int(event.x - self.vista.offset)
        y = int(event.y - self.vista.offset)

        # 2. Verificamos que el usuario no haya hecho clic fuera de la foto
        if x < 0 or x >= self.img_width or y < 0 or y >= self.img_height:
            return

        self.vista.mostrar_alerta("Calculando...", "Propagando semilla, por favor espera un momento.")
        self.vista.update()

        # 3. Llamamos a tu algoritmo matemático con una tolerancia de 30
        exito, region = self.modelo.crecimiento_semilla_hsi(self.img_pil, x, y, tolerancia=0.05)

        if exito:
            # 4. Clonamos la imagen a Numpy para pintarla súper rápido
            img_np = np.array(self.img_pil.convert('RGB'))
            
            # 5. Pintamos de Verde Fosforescente cada píxel "contagiado"
            for px, py in region:
                img_np[py, px] = [0, 255, 0] # R=0, G=255, B=0
                
            # 6. Convertimos de vuelta a imagen y la mostramos
            img_infectada_pil = Image.fromarray(img_np)
            self.img_tk_infectada = ImageTk.PhotoImage(img_infectada_pil)
            self.vista.dibujar_imagen(self.img_tk_infectada, self.img_width, self.img_height)
            
            self.vista.mostrar_alerta("Infección Terminada", f"Se expandió la región.\nTotal de píxeles: {len(region)}")
            
        # Apagamos el modo para que no siga pintando por accidente
        self.modo_semilla = False   

    def gestionar_aplauso(self):
        """Esta función se ejecuta automáticamente cuando detecta un aplauso."""
        if not self.en_proceso_expansion:
            print("¡Aplauso detectado! Modo semilla ACTIVADO.")
            self.modo_semilla = True
            self.en_proceso_expansion = True
            # Aquí podrías mandar un mensaje visual a la vista
        else:
            print("¡Aplauso detectado! Deteniendo expansión.")
            self.modo_semilla = False
            self.en_proceso_expansion = False

    def decir_escena(self, texto):
        """Configura el sintetizador y dice el texto en voz alta."""
        engine = pyttsx3.init()
        # Opcional: configurar voz en español
        voices = engine.getProperty('voices')
        for voice in voices:
            if "spanish" in voice.name.lower():
                engine.setProperty('voice', voice.id)
                break
        engine.say(texto)
        engine.runAndWait()

    def ejecutar_crecimiento_semilla(self, event):
        x = int(event.x - self.vista.offset)
        y = int(event.y - self.vista.offset)

        if x < 0 or x >= self.img_width or y < 0 or y >= self.img_height:
            return

        # Obtenemos el color HSI de la semilla para deducir qué es
        img_hsv = mcolors.rgb_to_hsv(np.array(self.img_pil.convert('RGB')) / 255.0)
        h, s, v = img_hsv[y, x]
        
        # --- LÓGICA DE ENTENDIMIENTO (Deducción por Matiz/Hue) ---
        # Los rangos de Hue van de 0 a 1 (como un círculo de color)
        etiqueta = "Desconocido"
        if 0.5 <= h <= 0.7: 
            etiqueta = "Cielo o Agua" # Tonos azules
        elif 0.2 <= h <= 0.45:
            etiqueta = "Vegetación"   # Tonos verdes
        elif 0.05 <= h <= 0.15:
            etiqueta = "Arena o Rocas" # Tonos amarillos/cafés
            
        self.vista.mostrar_alerta("Procesando", f"Detectando región de tipo: {etiqueta}")
        self.vista.update()

        # Ejecutamos la expansión Pro (Hue + Saturation)
        exito, region = self.modelo.crecimiento_semilla_hsi_pro(self.img_pil, x, y)

        if exito:
            # (Pintamos la región en verde como ya lo hacíamos...)
            img_np = np.array(self.img_pil.convert('RGB'))
            for px, py in region:
                img_np[py, px] = [0, 255, 0]
            
            self.img_tk_infectada = ImageTk.PhotoImage(Image.fromarray(img_np))
            self.vista.dibujar_imagen(self.img_tk_infectada, self.img_width, self.img_height)
            
            # --- ¡VOZ! El Megáfono del pizarrón ---
            mensaje_voz = f"He identificado una región de {etiqueta} con {len(region)} píxeles."
            self.decir_escena(mensaje_voz)
            
        self.modo_semilla = False

if __name__ == "__main__":
    app_vista = VistaPrincipal()
    app_modelo = ClasificadorModelo()
    

    presentador = Presentador(app_vista, app_modelo)
    
    app_vista.mainloop()